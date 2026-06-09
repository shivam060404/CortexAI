"""
FastAPI routes for CortexAI.
"""

from __future__ import annotations

import asyncio
import json
import os
import time
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException, WebSocket, WebSocketDisconnect
from langchain_core.messages import HumanMessage
from sqlalchemy import func, or_, select, update as sa_update, Integer

from backend.api.schemas import (
    ContextInjectRequest,
    CreateSessionRequest,
    ExecutionMetrics,
    SessionListResponse,
    SessionResponse,
    WatchSessionRequest,
    WorkspaceFileResponse,
)
from backend.auth.jwt_handler import verify_token
from backend.auth.models import User
from backend.auth.dependencies import get_current_active_user
from backend.core.alignment_engine import RESEARCH_MODES, align_query, needs_clarification
from backend.core.audit import audit_logger
from backend.core.graph import build_graph, cleanup_session, get_execution_metrics
from backend.core.guardrails import scan_llm_output, scan_user_input, verify_citations
from backend.core.logger import get_logger
from backend.core.preference_learning import get_user_preferences, learn_from_feedback
from backend.core.telemetry import trace_session, record_session_cost
from backend.core.report_sharing import create_share_link, get_shared_report, list_share_links, revoke_share_link
from backend.core.content_policy import check_content_policy
from backend.core.rate_limiter import check_websocket_rate_limit
from backend.core.scheduler import schedule_watch
from backend.core.session_store import session_store
from backend.core.supervisor_events import extract_supervisor_stream_messages
from backend.config import settings as _ws_settings
from backend.db.postgres import (
    AgentTrace,
    ExperimentTrack,
    FeedbackLog,
    KnowledgeEdge,
    KnowledgeNode,
    ResearchSession,
    SessionStatus,
    async_session,
)
from backend.db.tenant import bind_user_tenant_context, reset_tenant_context
from backend.db.workspace import WorkspaceManager
from backend.tools.memory_tools import update_user_memory
from backend.tools.planning_tools import get_cached_session_todos, get_session_todos, list_plan_versions, load_latest_plan

logger = get_logger(__name__)
router = APIRouter()
_workspace = WorkspaceManager()


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _owned_session_ids_query(current_user: User):
    user_uuid = uuid.UUID(str(current_user.id))
    return select(ResearchSession.id).where(ResearchSession.user_id == user_uuid)


async def _persist_session_to_db(session: dict) -> None:
    try:
        async with async_session() as db:
            existing = await db.get(ResearchSession, uuid.UUID(session["id"]))
            if existing is None:
                db.add(
                    ResearchSession(
                        id=uuid.UUID(session["id"]),
                        user_id=uuid.UUID(session["user_id"]),
                        title=session["title"],
                        user_request=session["user_request"],
                        status=SessionStatus(session["status"]),
                        final_report=session.get("final_report", ""),
                        iterations_used=session.get("iterations_used", 0),
                        tokens_used=session.get("tokens_used", 0),
                        tool_calls_count=session.get("tool_calls_count", 0),
                    )
                )
                await db.commit()
    except Exception as exc:
        logger.error("session_persist_failed", session_id=session["id"], error=str(exc))


async def _update_session_in_db(session_id: str, user_id: str, updates: dict) -> None:
    try:
        db_updates = dict(updates)
        if "status" in db_updates:
            db_updates["status"] = SessionStatus(db_updates["status"])
        db_updates["updated_at"] = datetime.now(timezone.utc)
        async with async_session() as db:
            await db.execute(
                sa_update(ResearchSession)
                .where(
                    ResearchSession.id == uuid.UUID(session_id),
                    ResearchSession.user_id == uuid.UUID(user_id),
                )
                .values(**db_updates)
            )
            await db.commit()
    except Exception as exc:
        logger.error("session_update_db_failed", session_id=session_id, error=str(exc))


async def _require_owned_session(session_id: str, current_user: User) -> dict:
    session = await session_store.get(session_id)
    if not session or session.get("user_id") != str(current_user.id):
        raise HTTPException(status_code=404, detail="Session not found")
    return session


async def _get_websocket_user(websocket: WebSocket) -> User:
    token = websocket.query_params.get("token")
    if not token:
        raise HTTPException(status_code=401, detail="Missing WebSocket token")

    try:
        payload = verify_token(token)
    except Exception as exc:
        raise HTTPException(status_code=401, detail="Invalid WebSocket token") from exc

    if payload.type != "access":
        raise HTTPException(status_code=401, detail="Invalid WebSocket token type")

    from backend.auth.dependencies import _resolve_user_from_credentials

    user = await _resolve_user_from_credentials(token=token, api_key=None)
    return user


@router.post("/api/context/pages")
async def inject_page_context(
    req: ContextInjectRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Receive webpage context injected from the Chrome extension."""
    try:
        session_id = req.session_id
        if session_id == "default":
            user_sessions = await session_store.list_by_user(str(current_user.id))
            active_sessions = sorted(
                [s for s in user_sessions if s["status"] in ("pending", "running")],
                key=lambda item: item["created_at"],
                reverse=True,
            )
            session_id = active_sessions[0]["id"] if active_sessions else "default"
        elif session_id != "default":
            await _require_owned_session(session_id, current_user)

        filename = f"context_{int(datetime.now().timestamp())}.json"
        content_dict = {
            "url": req.url,
            "title": req.title,
            "tags": req.tags,
            "note": req.note,
            "extracted_text": req.content,
            "injected_at": _now_iso(),
        }
        _workspace.write_file(session_id, filename, json.dumps(content_dict, indent=2))
        await update_user_memory(
            f"Saved page: {req.title or req.url}\nTags: {', '.join(req.tags)}\nNote: {req.note}",
            str(current_user.id),
        )
        return {"status": "success", "session_id": session_id, "file": filename}
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("context_injection_failed", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/api/sessions", response_model=SessionResponse)
async def create_session(
    req: CreateSessionRequest,
    current_user: User = Depends(get_current_active_user),
):
    """Create a new research session."""
    session_id = str(uuid.uuid4())
    session = {
        "id": session_id,
        "title": req.query[:100],
        "user_request": req.query,
        "status": "pending",
        "final_report": "",
        "iterations_used": 0,
        "tokens_used": 0,
        "tool_calls_count": 0,
        "user_id": str(current_user.id),
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
    }
    await session_store.set(session_id, session, user_id=str(current_user.id))
    await _persist_session_to_db(session)
    await audit_logger.log("session_create", user_id=str(current_user.id), details={"session_id": session_id})
    return SessionResponse(**session)


@router.get("/api/sessions", response_model=SessionListResponse)
async def list_sessions(current_user: User = Depends(get_current_active_user)):
    sessions = await session_store.list_by_user(str(current_user.id))
    return SessionListResponse(sessions=[SessionResponse(**s) for s in sessions], total=len(sessions))


@router.get("/api/sessions/{session_id}", response_model=SessionResponse)
async def get_session(session_id: str, current_user: User = Depends(get_current_active_user)):
    session = await _require_owned_session(session_id, current_user)
    return SessionResponse(**session)


@router.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    await session_store.delete(session_id, user_id=str(current_user.id))
    cleanup_session(session_id)

    from sqlalchemy import delete as sa_delete

    async with async_session() as db:
        await db.execute(
            sa_delete(ResearchSession).where(
                ResearchSession.id == uuid.UUID(session_id),
                ResearchSession.user_id == uuid.UUID(str(current_user.id)),
            )
        )
        await db.commit()

    await audit_logger.log("session_delete", user_id=str(current_user.id), details={"session_id": session_id})
    return {"status": "deleted"}


@router.post("/api/sessions/{session_id}/watch")
async def start_background_watch(
    session_id: str,
    req: WatchSessionRequest,
    current_user: User = Depends(get_current_active_user),
):
    await _require_owned_session(session_id, current_user)
    job_id = schedule_watch(
        session_id,
        str(current_user.id),
        req.topic,
        req.frequency_hours,
        organization_id=str(getattr(current_user, "organization_id", None) or current_user.id),
        role=("admin" if bool(getattr(current_user, "is_admin", False)) else getattr(current_user, "role", "owner")),
    )
    return {
        "status": "scheduled",
        "job_id": job_id,
        "topic": req.topic,
        "frequency_hours": req.frequency_hours,
    }


@router.get("/api/sessions/{session_id}/todos")
async def get_todos(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    return {"todos": await get_session_todos(session_id)}


@router.get("/api/sessions/{session_id}/plan")
async def get_current_plan(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    plan = await load_latest_plan(session_id)
    return {"plan": plan}


@router.get("/api/sessions/{session_id}/plans")
async def get_plan_versions(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    return {"plans": await list_plan_versions(session_id)}


@router.get("/api/sessions/{session_id}/files")
async def get_files(
    session_id: str,
    path: str = ".",
    current_user: User = Depends(get_current_active_user),
):
    await _require_owned_session(session_id, current_user)
    try:
        entries = _workspace.list_dir(session_id, path)
        return {"files": [WorkspaceFileResponse(**entry) for entry in entries]}
    except Exception as exc:
        return {"files": [], "error": str(exc)}


@router.get("/api/sessions/{session_id}/files/content")
async def get_file_content(
    session_id: str,
    path: str,
    current_user: User = Depends(get_current_active_user),
):
    await _require_owned_session(session_id, current_user)
    try:
        content = _workspace.read_file(session_id, path)
        return {"content": content, "path": path}
    except Exception as exc:
        raise HTTPException(status_code=404, detail=str(exc))


@router.get("/api/sessions/{session_id}/metrics")
async def get_metrics(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    metrics = get_execution_metrics(session_id)
    return ExecutionMetrics(**metrics) if metrics else ExecutionMetrics()


@router.get("/api/sessions/{session_id}/experiments")
async def get_experiments(session_id: str, current_user: User = Depends(get_current_active_user)):
    await _require_owned_session(session_id, current_user)
    try:
        async with async_session() as db:
            res = await db.execute(
                select(ExperimentTrack)
                .join(ResearchSession, ExperimentTrack.session_id == ResearchSession.id)
                .where(
                    ExperimentTrack.session_id == uuid.UUID(session_id),
                    ResearchSession.user_id == uuid.UUID(str(current_user.id)),
                )
                .order_by(ExperimentTrack.created_at.asc())
            )
            tracks = res.scalars().all()
            return {
                "experiments": [
                    {
                        "id": str(track.id),
                        "hypothesis": track.hypothesis,
                        "approach": track.approach,
                        "result": track.result,
                        "conclusion": track.conclusion,
                        "created_at": track.created_at.isoformat() if track.created_at else None,
                    }
                    for track in tracks
                ]
            }
    except Exception as exc:
        logger.error("get_experiments_error", error=str(exc))
        return {"experiments": [], "error": str(exc)}


@router.get("/api/sessions/{session_id}/traces")
async def get_session_traces(
    session_id: str,
    limit: int = 100,
    current_user: User = Depends(get_current_active_user),
):
    await _require_owned_session(session_id, current_user)
    try:
        async with async_session() as db:
            res = await db.execute(
                select(AgentTrace)
                .join(ResearchSession, AgentTrace.session_id == ResearchSession.id)
                .where(
                    AgentTrace.session_id == uuid.UUID(session_id),
                    ResearchSession.user_id == uuid.UUID(str(current_user.id)),
                )
                .order_by(AgentTrace.timestamp.asc())
                .limit(limit)
            )
            traces = res.scalars().all()
            return {
                "traces": [
                    {
                        "id": str(trace.id),
                        "event_type": trace.event_type,
                        "tool_name": trace.tool_name,
                        "input_data": trace.input_data,
                        "output_data": trace.output_data,
                        "latency_ms": round(trace.latency_ms, 1) if trace.latency_ms else 0,
                        "tokens_used": trace.tokens_used or 0,
                        "is_error": trace.is_error,
                        "error_detail": trace.error_detail or "",
                        "timestamp": trace.timestamp.isoformat() if trace.timestamp else None,
                    }
                    for trace in traces
                ]
            }
    except Exception as exc:
        logger.error("get_traces_error", error=str(exc))
        return {"traces": [], "error": str(exc)}


@router.get("/api/knowledge/nodes")
async def get_knowledge_nodes(limit: int = 50, current_user: User = Depends(get_current_active_user)):
    try:
        owned_session_ids = _owned_session_ids_query(current_user)
        async with async_session() as db:
            out_edges_sq = (
                select(KnowledgeEdge.source_id, func.count().label("cnt"))
                .where(KnowledgeEdge.session_id.in_(owned_session_ids))
                .group_by(KnowledgeEdge.source_id)
                .subquery()
            )
            in_edges_sq = (
                select(KnowledgeEdge.target_id, func.count().label("cnt"))
                .where(KnowledgeEdge.session_id.in_(owned_session_ids))
                .group_by(KnowledgeEdge.target_id)
                .subquery()
            )
            stmt = (
                select(
                    KnowledgeNode,
                    func.coalesce(out_edges_sq.c.cnt, 0).label("out_count"),
                    func.coalesce(in_edges_sq.c.cnt, 0).label("in_count"),
                )
                .outerjoin(out_edges_sq, KnowledgeNode.id == out_edges_sq.c.source_id)
                .outerjoin(in_edges_sq, KnowledgeNode.id == in_edges_sq.c.target_id)
                .where(KnowledgeNode.session_id.in_(owned_session_ids))
                .order_by(KnowledgeNode.created_at.desc())
                .limit(limit)
            )
            rows = (await db.execute(stmt)).all()
            return {
                "nodes": [
                    {
                        "id": str(node.id),
                        "name": node.name,
                        "node_type": node.node_type,
                        "attributes": node.attributes or {},
                        "edge_count": out_count + in_count,
                        "created_at": node.created_at.isoformat() if node.created_at else None,
                    }
                    for node, out_count, in_count in rows
                ],
                "total": len(rows),
            }
    except Exception as exc:
        logger.error("get_kg_nodes_error", error=str(exc))
        return {"nodes": [], "total": 0, "error": str(exc)}


@router.get("/api/knowledge/edges")
async def get_knowledge_edges(limit: int = 100, current_user: User = Depends(get_current_active_user)):
    try:
        owned_session_ids = _owned_session_ids_query(current_user)
        async with async_session() as db:
            edges = (
                await db.execute(
                    select(KnowledgeEdge)
                    .where(KnowledgeEdge.session_id.in_(owned_session_ids))
                    .order_by(KnowledgeEdge.created_at.desc())
                    .limit(limit)
                )
            ).scalars().all()
            edge_data = []
            for edge in edges:
                src_node = (await db.execute(select(KnowledgeNode).where(KnowledgeNode.id == edge.source_id))).scalar_one_or_none()
                tgt_node = (await db.execute(select(KnowledgeNode).where(KnowledgeNode.id == edge.target_id))).scalar_one_or_none()
                edge_data.append(
                    {
                        "id": str(edge.id),
                        "source": src_node.name if src_node else "unknown",
                        "target": tgt_node.name if tgt_node else "unknown",
                        "relation": edge.relation,
                        "source_id": str(edge.source_id),
                        "target_id": str(edge.target_id),
                    }
                )
            return {"edges": edge_data, "total": len(edge_data)}
    except Exception as exc:
        logger.error("get_kg_edges_error", error=str(exc))
        return {"edges": [], "total": 0, "error": str(exc)}


@router.get("/api/knowledge/search")
async def search_knowledge(q: str, limit: int = 20, current_user: User = Depends(get_current_active_user)):
    try:
        owned_session_ids = _owned_session_ids_query(current_user)
        async with async_session() as db:
            nodes = (
                await db.execute(
                    select(KnowledgeNode)
                    .where(
                        KnowledgeNode.session_id.in_(owned_session_ids),
                        KnowledgeNode.name.ilike(f"%{q}%"),
                    )
                    .limit(limit)
                )
            ).scalars().all()
            results = []
            for node in nodes:
                edges = (
                    await db.execute(
                        select(KnowledgeEdge)
                        .where(
                            KnowledgeEdge.session_id.in_(owned_session_ids),
                            or_(KnowledgeEdge.source_id == node.id, KnowledgeEdge.target_id == node.id),
                        )
                        .limit(10)
                    )
                ).scalars().all()
                relations = []
                for edge in edges:
                    other_id = edge.target_id if edge.source_id == node.id else edge.source_id
                    other_node = (
                        await db.execute(select(KnowledgeNode).where(KnowledgeNode.id == other_id))
                    ).scalar_one_or_none()
                    if other_node:
                        direction = "->" if edge.source_id == node.id else "<-"
                        relations.append(f"{direction} [{edge.relation}] {other_node.name}")
                results.append(
                    {
                        "id": str(node.id),
                        "name": node.name,
                        "node_type": node.node_type,
                        "relations": relations,
                    }
                )
            return {"results": results}
    except Exception as exc:
        logger.error("kg_search_error", error=str(exc))
        return {"results": [], "error": str(exc)}


@router.get("/api/experiments/stats")
async def get_experiment_stats(current_user: User = Depends(get_current_active_user)):
    try:
        user_uuid = uuid.UUID(str(current_user.id))
        async with async_session() as db:
            total_count = (
                await db.execute(
                    select(func.count())
                    .select_from(ExperimentTrack)
                    .join(ResearchSession, ExperimentTrack.session_id == ResearchSession.id)
                    .where(ResearchSession.user_id == user_uuid)
                )
            ).scalar() or 0
            session_count = (
                await db.execute(
                    select(func.count(func.distinct(ExperimentTrack.session_id)))
                    .join(ResearchSession, ExperimentTrack.session_id == ResearchSession.id)
                    .where(ResearchSession.user_id == user_uuid)
                )
            ).scalar() or 0
            recent = (
                await db.execute(
                    select(ExperimentTrack)
                    .join(ResearchSession, ExperimentTrack.session_id == ResearchSession.id)
                    .where(ResearchSession.user_id == user_uuid)
                    .order_by(ExperimentTrack.created_at.desc())
                    .limit(5)
                )
            ).scalars().all()
            return {
                "total_experiments": total_count,
                "sessions_with_experiments": session_count,
                "recent": [
                    {
                        "id": str(track.id),
                        "hypothesis": track.hypothesis[:100],
                        "conclusion": (track.conclusion or "")[:100],
                        "session_id": str(track.session_id),
                        "created_at": track.created_at.isoformat() if track.created_at else None,
                    }
                    for track in recent
                ],
            }
    except Exception as exc:
        logger.error("experiment_stats_error", error=str(exc))
        return {"total_experiments": 0, "sessions_with_experiments": 0, "recent": [], "error": str(exc)}


@router.get("/api/research/modes")
async def get_research_modes(current_user: User = Depends(get_current_active_user)):
    return {
        "modes": [
            {"id": key, "label": value["label"], "depth": value["depth"], "max_sources": value["max_sources"]}
            for key, value in RESEARCH_MODES.items()
        ]
    }


@router.post("/api/research/align")
async def align_research_query(req: dict, current_user: User = Depends(get_current_active_user)):
    query = req.get("query", "")
    mode = req.get("mode", "deep")
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query is required")
    clarification = needs_clarification(query)
    if clarification:
        return {"needs_clarification": True, **clarification}
    prefs = await get_user_preferences(str(current_user.id))
    alignment = await align_query(query, mode=mode, user_prefs=prefs)
    return {
        "needs_clarification": False,
        "original_query": query,
        "refined_query": alignment["refined_query"],
        "sub_queries": alignment["sub_queries"],
        "search_strategy": alignment["search_strategy"],
        "mode": mode,
        "mode_config": alignment["mode_config"],
    }


@router.get("/api/sessions/{session_id}/export/{fmt}")
async def export_session_report(
    session_id: str,
    fmt: str,
    current_user: User = Depends(get_current_active_user),
):
    """Export a session report as PDF or DOCX (Feature Gap #4)."""
    from fastapi.responses import FileResponse

    await _require_owned_session(session_id, current_user)
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Supported formats: pdf, docx")

    try:
        report_md = _workspace.read_file(session_id, "report.md")
    except Exception:
        raise HTTPException(status_code=404, detail="No report found for this session")

    title = session_id[:8]
    session = await session_store.get(session_id)
    if session:
        title = session.get("title", session_id[:8])

    if fmt == "pdf":
        filename = f"report_{session_id[:8]}.pdf"
        try:
            from weasyprint import HTML
            from backend.tools.export_tools import _markdown_to_html
            html_body = _markdown_to_html(report_md, title)
            html_doc = f"""<!DOCTYPE html><html><head><meta charset='utf-8'><style>
body{{font-family:Arial,sans-serif;margin:40px;line-height:1.6}}
h1{{color:#1a1a2e}}h2{{color:#16213e}}h3{{color:#0f3460}}
</style></head><body>{html_body}</body></html>"""
            workspace_dir = _workspace.get_workspace_path(session_id)
            file_path = os.path.join(workspace_dir, filename)
            HTML(string=html_doc).write_pdf(file_path)
            return FileResponse(file_path, media_type="application/pdf", filename=filename)
        except ImportError:
            raise HTTPException(status_code=501, detail="PDF export not available on this server")
    else:
        filename = f"report_{session_id[:8]}.docx"
        try:
            from docx import Document
            from docx.shared import Pt as DocxPt
            doc = Document()
            doc.add_heading(title, level=0)
            for line in report_md.split("\n"):
                s = line.strip()
                if not s:
                    continue
                elif s.startswith("## "):
                    doc.add_heading(s[3:].replace("**", "").strip(), level=2)
                elif s.startswith("### "):
                    doc.add_heading(s[4:].replace("**", "").strip(), level=3)
                elif s.startswith(("- ", "* ")):
                    doc.add_paragraph(s[2:].replace("**", "").strip(), style="List Bullet")
                else:
                    doc.add_paragraph(s.replace("**", "").strip())
            workspace_dir = _workspace.get_workspace_path(session_id)
            file_path = os.path.join(workspace_dir, filename)
            doc.save(file_path)
            return FileResponse(
                file_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=filename,
            )
        except ImportError:
            raise HTTPException(status_code=501, detail="DOCX export not available on this server")


@router.post("/api/feedback")
async def submit_feedback(req: dict, current_user: User = Depends(get_current_active_user)):
    session_id = req.get("session_id", "")
    rating = req.get("rating", 0)
    comment = req.get("comment", "")
    if not session_id or rating is None:
        raise HTTPException(status_code=400, detail="session_id and rating required")
    session = await _require_owned_session(session_id, current_user)
    try:
        async with async_session() as db:
            db.add(
                FeedbackLog(
                    user_id=uuid.UUID(str(current_user.id)),
                    session_id=uuid.UUID(session_id),
                    rating=rating,
                    comment=comment,
                    query=session.get("user_request", ""),
                    research_mode=req.get("mode", "deep"),
                )
            )
            await db.commit()
        await learn_from_feedback(session_id, rating, comment, str(current_user.id))
        await audit_logger.log(
            "feedback_submit",
            user_id=str(current_user.id),
            details={"session_id": session_id, "rating": rating},
        )
        return {"status": "ok", "message": "Thank you for your feedback"}
    except Exception as exc:
        logger.error("feedback_error", error=str(exc))
        return {"status": "error", "message": str(exc)}


@router.get("/api/preferences")
async def get_preferences(current_user: User = Depends(get_current_active_user)):
    prefs = await get_user_preferences(str(current_user.id))
    return {"preferences": prefs}


# ─── Analytics Dashboard API (Task 14c) ───


@router.get("/api/analytics/usage")
async def analytics_usage(
    days: int = 30,
    current_user: User = Depends(get_current_active_user),
):
    """Session count, tokens used, and estimated cost over time."""
    from datetime import timedelta
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    user_uuid = uuid.UUID(str(current_user.id))

    try:
        async with async_session() as db:
            res = await db.execute(
                select(
                    func.count().label("total_sessions"),
                    func.coalesce(func.sum(ResearchSession.tokens_used), 0).label("total_tokens"),
                    func.coalesce(func.sum(ResearchSession.tool_calls_count), 0).label("total_tool_calls"),
                    func.coalesce(func.sum(ResearchSession.iterations_used), 0).label("total_iterations"),
                ).where(
                    ResearchSession.user_id == user_uuid,
                    ResearchSession.created_at >= cutoff,
                )
            )
            row = res.one()

        total_tokens = int(row.total_tokens or 0)
        # Estimate cost from settings
        from backend.config import settings as _s
        estimated_cost = (
            (total_tokens * 0.5 * _s.COST_PER_1M_INPUT_TOKENS / 1_000_000) +
            (total_tokens * 0.5 * _s.COST_PER_1M_OUTPUT_TOKENS / 1_000_000)
        )

        return {
            "period_days": days,
            "total_sessions": int(row.total_sessions or 0),
            "total_tokens": total_tokens,
            "total_tool_calls": int(row.total_tool_calls or 0),
            "total_iterations": int(row.total_iterations or 0),
            "estimated_cost_usd": round(estimated_cost, 4),
        }
    except Exception as e:
        logger.error("analytics_usage_error", error=str(e))
        return {"period_days": days, "total_sessions": 0, "total_tokens": 0,
                "total_tool_calls": 0, "total_iterations": 0, "estimated_cost_usd": 0}


@router.get("/api/analytics/costs")
async def analytics_costs(
    days: int = 30,
    current_user: User = Depends(get_current_active_user),
):
    """Per-session cost breakdown."""
    from datetime import timedelta
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    user_uuid = uuid.UUID(str(current_user.id))

    try:
        from backend.config import settings as _s
        async with async_session() as db:
            res = await db.execute(
                select(ResearchSession)
                .where(
                    ResearchSession.user_id == user_uuid,
                    ResearchSession.created_at >= cutoff,
                )
                .order_by(ResearchSession.created_at.desc())
                .limit(50)
            )
            sessions = res.scalars().all()

        breakdown = []
        for s in sessions:
            tokens = s.tokens_used or 0
            cost = (
                (tokens * 0.5 * _s.COST_PER_1M_INPUT_TOKENS / 1_000_000) +
                (tokens * 0.5 * _s.COST_PER_1M_OUTPUT_TOKENS / 1_000_000)
            )
            breakdown.append({
                "session_id": str(s.id),
                "title": s.title[:60],
                "tokens_used": tokens,
                "iterations": s.iterations_used or 0,
                "tool_calls": s.tool_calls_count or 0,
                "estimated_cost_usd": round(cost, 4),
                "status": s.status.value if hasattr(s.status, "value") else str(s.status),
                "created_at": s.created_at.isoformat() if s.created_at else None,
            })

        return {"period_days": days, "sessions": breakdown}
    except Exception as e:
        logger.error("analytics_costs_error", error=str(e))
        return {"period_days": days, "sessions": []}


@router.get("/api/analytics/performance")
async def analytics_performance(
    days: int = 30,
    current_user: User = Depends(get_current_active_user),
):
    """Average latency, success rates, and iteration statistics."""
    from datetime import timedelta
    cutoff = datetime.now(timezone.utc) - timedelta(days=days)
    user_uuid = uuid.UUID(str(current_user.id))

    try:
        async with async_session() as db:
            # Session status distribution
            status_res = await db.execute(
                select(
                    ResearchSession.status,
                    func.count().label("count"),
                    func.coalesce(func.avg(ResearchSession.iterations_used), 0).label("avg_iterations"),
                    func.coalesce(func.avg(ResearchSession.tokens_used), 0).label("avg_tokens"),
                )
                .where(
                    ResearchSession.user_id == user_uuid,
                    ResearchSession.created_at >= cutoff,
                )
                .group_by(ResearchSession.status)
            )
            status_rows = status_res.all()

            # Agent trace latency stats
            owned_ids = _owned_session_ids_query(current_user)
            trace_res = await db.execute(
                select(
                    func.count().label("total_traces"),
                    func.coalesce(func.avg(AgentTrace.latency_ms), 0).label("avg_latency_ms"),
                    func.coalesce(func.sum(func.cast(AgentTrace.is_error, Integer)), 0).label("error_count"),
                )
                .where(AgentTrace.session_id.in_(owned_ids))
            )
            trace_row = trace_res.one()

        status_dist = {}
        total_sessions = 0
        for row in status_rows:
            status_val = row.status.value if hasattr(row.status, "value") else str(row.status)
            status_dist[status_val] = {
                "count": int(row.count),
                "avg_iterations": round(float(row.avg_iterations), 1),
                "avg_tokens": int(row.avg_tokens or 0),
            }
            total_sessions += int(row.count)

        total_traces = int(trace_row.total_traces or 0)
        error_count = int(trace_row.error_count or 0)
        success_rate = ((total_traces - error_count) / max(total_traces, 1)) * 100

        return {
            "period_days": days,
            "total_sessions": total_sessions,
            "status_distribution": status_dist,
            "total_tool_calls": total_traces,
            "avg_tool_latency_ms": round(float(trace_row.avg_latency_ms or 0), 1),
            "tool_error_count": error_count,
            "tool_success_rate": round(success_rate, 1),
        }
    except Exception as e:
        logger.error("analytics_performance_error", error=str(e))
        return {"period_days": days, "total_sessions": 0, "status_distribution": {},
                "total_tool_calls": 0, "avg_tool_latency_ms": 0, "tool_error_count": 0, "tool_success_rate": 0}


# ---------------------------------------------------------------------------
# Report Sharing (Feature Gap #5 / Task 17)
# ---------------------------------------------------------------------------
@router.post("/api/sessions/{session_id}/share")
async def create_share_endpoint(
    session_id: str,
    is_public: bool = False,
    expires_days: int = 30,
    user: User = Depends(get_current_active_user),
):
    """Create a shareable link for a completed research report."""
    try:
        result = await create_share_link(
            session_id=session_id,
            user_id=str(user.id),
            is_public=is_public,
            expires_days=expires_days,
        )
        await audit_logger.log("share_link_created", user_id=str(user.id),
                               details={"session_id": session_id, "is_public": is_public})
        return result
    except Exception as e:
        logger.error("create_share_link_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/sessions/{session_id}/shares")
async def list_shares_endpoint(
    session_id: str,
    user: User = Depends(get_current_active_user),
):
    """List all share links for a session."""
    try:
        return await list_share_links(session_id=session_id, user_id=str(user.id))
    except Exception as e:
        logger.error("list_share_links_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/api/shared/{token}")
async def revoke_share_endpoint(
    token: str,
    user: User = Depends(get_current_active_user),
):
    """Revoke (delete) a share link."""
    try:
        result = await revoke_share_link(token=token, user_id=str(user.id))
        await audit_logger.log("share_link_revoked", user_id=str(user.id),
                               details={"token_prefix": token[:8]})
        return result
    except Exception as e:
        logger.error("revoke_share_link_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/shared/{token}")
async def view_shared_report(token: str):
    """View a shared report (no auth required for public shares)."""
    try:
        report = await get_shared_report(token)
        if not report:
            raise HTTPException(status_code=404, detail="Share link not found or expired")
        return report
    except HTTPException:
        raise
    except Exception as e:
        logger.error("view_shared_report_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# SOC2 Compliance Audit Endpoints (Feature Gap #12 / Task 18)
# ---------------------------------------------------------------------------
@router.get("/api/audit/summary")
async def audit_summary_endpoint(
    days: int = 30,
    user: User = Depends(get_current_active_user),
):
    """Get aggregate audit statistics for the given period."""
    try:
        return await audit_logger.get_audit_summary(days=days)
    except Exception as e:
        logger.error("audit_summary_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/audit/export")
async def audit_export_endpoint(
    days: int = 30,
    event_type: str | None = None,
    format: str = "json",
    user: User = Depends(get_current_active_user),
):
    """Export audit logs for compliance review (JSON or CSV)."""
    try:
        from datetime import timedelta
        end_date = datetime.now(timezone.utc)
        start_date = end_date - timedelta(days=days)
        result = await audit_logger.export_logs(
            start_date=start_date,
            end_date=end_date,
            event_type=event_type,
            format=format,
        )
        if format == "csv":
            from fastapi.responses import PlainTextResponse
            return PlainTextResponse(content=result, media_type="text/csv")
        return result
    except Exception as e:
        logger.error("audit_export_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/audit/retention")
async def audit_retention_endpoint(
    user: User = Depends(get_current_active_user),
):
    """Apply audit log retention policy (delete logs older than AUDIT_RETENTION_DAYS)."""
    try:
        return await audit_logger.apply_retention_policy()
    except Exception as e:
        logger.error("audit_retention_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Tool Registry & MCP Hot-Reload (Feature Gap #10 / Task 19)
# ---------------------------------------------------------------------------
@router.get("/api/tools")
async def list_tools_endpoint(
    source: str | None = None,
    category: str | None = None,
    user: User = Depends(get_current_active_user),
):
    """List all registered tools with optional filtering."""
    from backend.core.tool_registry import tool_registry
    if source:
        tools = tool_registry.get_tools_by_source(source)
    elif category:
        tools = tool_registry.get_tools_by_category(category)
    else:
        tools = tool_registry.get_enabled_tools()
    return {
        "total": len(tools),
        "tools": [
            {"id": t.tool_id, "name": t.name, "source": t.source,
             "description": t.description, "categories": t.categories,
             "enabled": t.is_enabled}
            for t in tools
        ],
    }


@router.get("/api/tools/status")
async def tool_registry_status(
    user: User = Depends(get_current_active_user),
):
    """Get tool registry status for observability."""
    from backend.core.tool_registry import tool_registry
    return tool_registry.get_status()


@router.post("/api/mcp/reload")
async def mcp_hot_reload_endpoint(
    user: User = Depends(get_current_active_user),
):
    """Hot-reload MCP server configuration from disk."""
    try:
        from backend.mcp.server_registry import MCPServerRegistry
        mcp_reg = MCPServerRegistry()
        result = await mcp_reg.hot_reload_config()
        await audit_logger.log_config_change(
            user_id=str(user.id),
            setting_name="mcp_servers_config",
            old_value="previous",
            new_value="reloaded",
        )
        return result
    except Exception as e:
        logger.error("mcp_hot_reload_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Real-Time Collaboration (Feature Gap #5 / Task 20)
# ---------------------------------------------------------------------------
@router.post("/api/sessions/{session_id}/collaborate/join")
async def join_session_endpoint(
    session_id: str,
    role: str = "viewer",
    user: User = Depends(get_current_active_user),
):
    """Join a research session as a collaborator."""
    from backend.core.collaboration import collaboration_manager
    try:
        result = await collaboration_manager.join_session(
            session_id=session_id, user_id=str(user.id), role=role,
        )
        await audit_logger.log("session_collaborate_join", user_id=str(user.id),
                               details={"session_id": session_id, "role": role})
        # Broadcast join event to existing participants
        await collaboration_manager.broadcast_event(
            session_id, "participant_joined",
            {"user_id": str(user.id), "role": role},
            exclude_user=str(user.id),
        )
        return result
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        logger.error("collaboration_join_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/api/sessions/{session_id}/collaborate/leave")
async def leave_session_endpoint(
    session_id: str,
    user: User = Depends(get_current_active_user),
):
    """Leave a research session."""
    from backend.core.collaboration import collaboration_manager
    try:
        result = await collaboration_manager.leave_session(
            session_id=session_id, user_id=str(user.id),
        )
        await collaboration_manager.broadcast_event(
            session_id, "participant_left",
            {"user_id": str(user.id)},
            exclude_user=str(user.id),
        )
        return result
    except Exception as e:
        logger.error("collaboration_leave_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/sessions/{session_id}/collaborate/participants")
async def list_participants_endpoint(
    session_id: str,
    user: User = Depends(get_current_active_user),
):
    """List all participants in a research session."""
    from backend.core.collaboration import collaboration_manager
    try:
        return await collaboration_manager.list_participants(session_id=session_id)
    except Exception as e:
        logger.error("collaboration_list_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/api/sessions/{session_id}/collaborate/role")
async def update_role_endpoint(
    session_id: str,
    target_user_id: str,
    new_role: str,
    user: User = Depends(get_current_active_user),
):
    """Update a participant's role (requires owner/admin)."""
    from backend.core.collaboration import collaboration_manager
    try:
        result = await collaboration_manager.update_role(
            session_id=session_id, user_id=target_user_id, new_role=new_role,
        )
        await collaboration_manager.broadcast_event(
            session_id, "role_updated",
            {"user_id": target_user_id, "new_role": new_role},
        )
        return result
    except Exception as e:
        logger.error("collaboration_role_update_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/api/collaboration/active")
async def active_collaborations_endpoint(
    user: User = Depends(get_current_active_user),
):
    """Get sessions with active real-time connections."""
    from backend.core.collaboration import collaboration_manager
    return collaboration_manager.get_active_sessions()


# ---------------------------------------------------------------------------
# Agent Harness Configuration & Health (Task 21)
# ---------------------------------------------------------------------------
@router.get("/api/harness/config")
async def harness_config_endpoint(
    org_id: str | None = None,
    user: User = Depends(get_current_active_user),
):
    """Get the full Agent Harness configuration."""
    from backend.core.harness import agent_harness
    return agent_harness.get_config(org_id=org_id)


@router.get("/api/harness/health")
async def harness_health_endpoint(
    user: User = Depends(get_current_active_user),
):
    """Run a live health check across all Agent Harness pillars."""
    from backend.core.harness import agent_harness
    return await agent_harness.run_health_check()


@router.get("/api/harness/score")
async def harness_score_endpoint(
    user: User = Depends(get_current_active_user),
):
    """Get detailed Agent Harness pillar scoring."""
    from backend.core.harness import agent_harness
    return agent_harness.get_total_harness_score()


# ─── Search Orchestration Endpoints (Phase 3) ───


@router.post("/api/search/batch")
async def batch_search(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Execute batch search with multiple queries in parallel."""
    from backend.core.search_orchestrator import search_orchestrator, SearchQuery, SearchDepth, SearchPriority

    queries_raw = req.get("queries", [])
    if not queries_raw:
        raise HTTPException(status_code=400, detail="queries array is required")

    queries = []
    for q in queries_raw:
        if isinstance(q, str):
            queries.append(SearchQuery(text=q))
        elif isinstance(q, dict):
            queries.append(SearchQuery(
                text=q.get("text", q.get("query", "")),
                priority=SearchPriority(q.get("priority", "normal")),
                depth=SearchDepth(q.get("depth", "standard")),
                category=q.get("category", "web"),
            ))

    max_concurrent = req.get("max_concurrent", 5)
    result = await search_orchestrator.execute_batch(queries, max_concurrent=max_concurrent)
    return {
        "total_queries": result.total_queries,
        "successful": result.successful_queries,
        "failed": result.failed_queries,
        "unique_urls": result.unique_urls,
        "duration_ms": result.total_duration_ms,
        "results": [
            {"url": r.url, "title": r.title, "content": r.content[:500],
             "score": r.score, "source": r.source}
            for r in result.results[:50]
        ],
    }


@router.post("/api/search/deep")
async def deep_search(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Execute iterative deep search with automatic follow-up generation."""
    from backend.core.search_orchestrator import search_orchestrator

    query = req.get("query", "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    max_depth = min(req.get("max_depth", 3), 5)
    follow_ups = min(req.get("follow_ups_per_level", 3), 5)
    category = req.get("category", "web")

    result = await search_orchestrator.execute_deep(
        initial_query=query, max_depth=max_depth,
        follow_ups_per_level=follow_ups, category=category,
    )
    return {
        "total_queries": result.total_queries,
        "successful": result.successful_queries,
        "unique_urls": result.unique_urls,
        "duration_ms": result.total_duration_ms,
        "results": [
            {"url": r.url, "title": r.title, "content": r.content[:500],
             "score": r.score, "source": r.source, "depth": r.depth_level}
            for r in result.results[:100]
        ],
    }


# ─── Predictive Research Planning Endpoints (Phase 4) ───


@router.post("/api/planning/predict")
async def predict_research_topics(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Predict next research topics based on session history."""
    from backend.core.predictive_planner import predictive_planner

    session_history = req.get("session_history", [])
    result = await predictive_planner.predict_next_topics(str(user.id), session_history)
    return {
        "predicted_topics": [
            {"topic": t.topic, "confidence": t.confidence, "related": t.related_topics}
            for t in result.predicted_topics
        ],
        "knowledge_gaps": [
            {"topic": g.topic, "reason": g.reason, "priority": g.priority, "suggested_queries": g.suggested_queries}
            for g in result.knowledge_gaps
        ],
        "research_score": result.research_score,
        "insights": result.insights,
    }


@router.post("/api/planning/generate")
async def generate_research_plan(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Generate a structured research plan for a topic."""
    from backend.core.predictive_planner import predictive_planner

    topic = req.get("topic", "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="topic is required")

    depth = req.get("depth", "default")
    plan = await predictive_planner.generate_research_plan(topic, depth=depth)
    return {
        "title": plan.title,
        "description": plan.description,
        "estimated_time": plan.estimated_total_time,
        "difficulty": plan.difficulty,
        "prerequisites": plan.prerequisites,
        "milestones": [
            {"title": m.title, "description": m.description, "time": m.estimated_time,
             "queries": m.queries, "order": m.order, "dependencies": m.dependencies}
            for m in plan.milestones
        ],
    }


@router.post("/api/planning/suggest")
async def suggest_research_plans(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Suggest research plans based on predicted topics."""
    from backend.core.predictive_planner import predictive_planner

    session_history = req.get("session_history", [])
    plans = await predictive_planner.suggest_research_plans(str(user.id), session_history)
    return {
        "plans": [
            {"title": p.title, "description": p.description, "estimated_time": p.estimated_total_time,
             "difficulty": p.difficulty, "milestone_count": len(p.milestones),
             "milestones": [{"title": m.title, "description": m.description}
                            for m in p.milestones]}
            for p in plans
        ],
    }


# ─── Collected Pages Endpoints (Phase 3 — Web Compare) ───

# In-memory store for collected pages (per-user, backed by Redis in production)
_collected_pages: dict[str, list[dict]] = {}


@router.post("/api/context/pages")
async def collect_pages(req: dict, user: User = Depends(get_current_active_user)):
    """Collect page data from Chrome extension for comparison."""
    uid = str(user.id)
    pages = req.get("pages", [])
    action = req.get("action", "collect")

    if uid not in _collected_pages:
        _collected_pages[uid] = []

    if action == "collect":
        for page in pages:
            # Deduplicate by URL
            _collected_pages[uid] = [p for p in _collected_pages[uid] if p.get("url") != page.get("url")]
            _collected_pages[uid].insert(0, {**page, "collected_at": time.time()})
        # Cap at 50 pages
        _collected_pages[uid] = _collected_pages[uid][:50]
        return {"collected": len(pages), "total": len(_collected_pages[uid])}

    elif action == "compare":
        comparison = req.get("comparison", {})
        return {"received": True, "pages": len(pages), "comparison": comparison}

    return {"status": "ok"}


@router.get("/api/context/collected-pages")
async def get_collected_pages(user: User = Depends(get_current_active_user)):
    """Get all collected pages for the current user."""
    uid = str(user.id)
    return {"pages": _collected_pages.get(uid, [])}


@router.post("/api/context/compare")
async def compare_collected_pages(
    req: dict,
    user: User = Depends(get_current_active_user),
):
    """Compare selected pages and generate analysis."""
    uid = str(user.id)
    urls = req.get("urls", [])
    user_pages = _collected_pages.get(uid, [])

    selected = [p for p in user_pages if p.get("url") in urls] if urls else user_pages[:5]

    if len(selected) < 2:
        raise HTTPException(status_code=400, detail="Need at least 2 pages to compare")

    # Generate comparison
    all_tags = {}
    for p in selected:
        for tag in p.get("tags", []):
            all_tags[tag] = all_tags.get(tag, 0) + 1

    common_tags = [{"tag": t, "count": c} for t, c in all_tags.items() if c >= 2]
    unique_tags = [t for t, c in all_tags.items() if c == 1]

    return {
        "pages": selected,
        "common_tags": common_tags,
        "unique_tags": unique_tags,
        "summary": {
            "page_count": len(selected),
            "total_word_count": sum(p.get("wordCount", 0) for p in selected),
        },
    }


@router.websocket("/ws/{session_id}")
async def websocket_research(websocket: WebSocket, session_id: str):
    user: User | None = None
    listener_task: asyncio.Task | None = None
    heartbeat_task: asyncio.Task | None = None
    tenant_token = None

    try:
        user = await _get_websocket_user(websocket)
        tenant_token = bind_user_tenant_context(user, source="websocket")
        is_allowed, headers = await check_websocket_rate_limit(
            client_ip=websocket.client.host if websocket.client else "unknown",
            user_id=str(user.id),
        )
        if not is_allowed:
            await websocket.close(code=4408, reason=headers.get("Retry-After", "Rate limit exceeded"))
            return

        await websocket.accept()

        # --- WebSocket Heartbeat (Arch Issue #8) ---
        _last_pong: dict[str, float] = {"ts": asyncio.get_event_loop().time()}

        async def _heartbeat():
            """Send periodic pings; close if no pong within timeout."""
            interval = _ws_settings.WS_HEARTBEAT_INTERVAL
            timeout = _ws_settings.WS_HEARTBEAT_TIMEOUT
            try:
                while True:
                    await asyncio.sleep(interval)
                    try:
                        await websocket.send_json({"type": "ping"})
                    except Exception:
                        break
                    elapsed = asyncio.get_event_loop().time() - _last_pong["ts"]
                    if elapsed > timeout:
                        logger.warning("ws_heartbeat_timeout", session_id=session_id)
                        try:
                            await websocket.close(code=4408, reason="Heartbeat timeout")
                        except Exception:
                            pass
                        break
            except asyncio.CancelledError:
                pass

        heartbeat_task = asyncio.create_task(_heartbeat())

        await audit_logger.log("ws_connect", user_id=str(user.id), details={"session_id": session_id})

        existing_session = await session_store.get(session_id)
        if existing_session and existing_session.get("user_id") != str(user.id):
            await websocket.send_json({"type": "error", "data": {"message": "Session not found"}})
            await websocket.close(code=4404)
            return

        data = await websocket.receive_text()
        try:
            payload = json.loads(data) if data.lstrip().startswith("{") else {"query": data}
        except json.JSONDecodeError:
            await websocket.send_json({"type": "error", "data": {"message": "Invalid JSON payload"}})
            await websocket.close(code=4400)
            return
        query = payload.get("query", data)
        research_mode = payload.get("mode", "deep")
        if not str(query).strip():
            await websocket.send_json({"type": "error", "data": {"message": "Query is required"}})
            await websocket.close(code=4400)
            return

        input_scan = scan_user_input(query)
        if not input_scan.is_safe:
            await audit_logger.log(
                "auth_failure",
                user_id=str(user.id),
                details={"session_id": session_id, "reason": "unsafe_query"},
            )
            await websocket.send_json(
                {
                    "type": "error",
                    "data": {
                        "message": input_scan.rejection_message,
                        "blocked": True,
                        "risk_score": input_scan.risk_score,
                    },
                }
            )
            await websocket.close(code=4400)
            return

        query = input_scan.sanitized_query or query
        session = existing_session or {
            "id": session_id,
            "title": query[:100],
            "user_request": query,
            "status": "pending",
            "final_report": "",
            "iterations_used": 0,
            "tokens_used": 0,
            "tool_calls_count": 0,
            "user_id": str(user.id),
            "created_at": _now_iso(),
            "updated_at": _now_iso(),
        }
        if existing_session is None:
            await session_store.set(session_id, session, user_id=str(user.id))
            await _persist_session_to_db(session)

        session = await session_store.update(session_id, {"status": "running", "user_request": query}) or session
        await _update_session_in_db(session_id, str(user.id), {"status": "running", "user_request": query})

        await websocket.send_json({"type": "status", "data": {"status": "running", "message": "Research started"}})
        await websocket.send_json({"type": "thinking", "data": {"message": "Aligning query with user preferences..."}})

        user_prefs = await get_user_preferences(str(user.id))
        alignment = await align_query(query, mode=research_mode, user_prefs=user_prefs)
        aligned_query = alignment["refined_query"]
        sub_queries = alignment.get("sub_queries", [])
        mode_config = alignment.get("mode_config", {})
        strategy = alignment.get("search_strategy", "balanced")

        await websocket.send_json(
            {
                "type": "thinking",
                "data": {
                    "message": f"Aligned query using {strategy} strategy in {research_mode} mode"
                },
            }
        )

        enriched_query = f"""{aligned_query}

[ALIGNMENT CONTEXT]
- Research Mode: {mode_config.get('label', 'Deep')} ({mode_config.get('depth', 'comprehensive')})
- Mode Instructions: {mode_config.get('instructions', '')}
- Search Strategy: {strategy}
- Sub-questions to address: {chr(10).join(f'  {i + 1}. {item}' for i, item in enumerate(sub_queries)) if sub_queries else 'None'}
- User preferences: {user_prefs if user_prefs else 'No preferences learned yet'}
"""
        await update_user_memory(query, str(user.id))
        graph = await build_graph(session_id)

        initial_state = {
            "messages": [HumanMessage(content=enriched_query)],
            "session_id": session_id,
            "status": "running",
            "supervisor_phase": "RECEIVED",
            "analysis_summary": "",
            "evidence_confidence": 0.0,
            "needs_replan": False,
            "final_output": "",
            "supervisor_events": [],
            "iteration": 0,
            "consecutive_failures": 0,
            "accessed_urls": set(),
            "hitl_mode": "supervised",
            "pending_approval": None,
            "user_modifications": [],
            "next": "Received",
            "plan_id": None,
            "plan_version": 0,
            "plan_status": "missing",
            "plan_summary": "",
            "last_worker": None,
        }

        from backend.core.hitl import HITLManager

        async def listen_to_client():
            while True:
                try:
                    msg = await websocket.receive_text()
                    client_payload = json.loads(msg)
                    if client_payload.get("type") == "pong":
                        _last_pong["ts"] = asyncio.get_event_loop().time()
                    elif client_payload.get("type") == "hitl_resume":
                        action = client_payload.get("data", {}).get("action", "continue")
                        modifications = client_payload.get("data", {}).get("modifications", {})
                        modifications["action"] = action
                        HITLManager.resume_with_input(session_id, modifications)
                except WebSocketDisconnect:
                    break
                except Exception as exc:
                    logger.error("ws_listen_error", error=str(exc))
                    break

        listener_task = asyncio.create_task(listen_to_client())
        tracked_urls = set()
        last_supervisor_phase = None
        sent_supervisor_event_count = 0

        async for event in graph.astream_events(
            initial_state,
            {"configurable": {"thread_id": session_id}},
            version="v2",
        ):
            try:
                event_type = event.get("event", "")
                event_name = event.get("name", "")
                output = event.get("data", {}).get("output", {})
                if isinstance(output, dict) and "accessed_urls" in output:
                    tracked_urls.update(output["accessed_urls"])

                supervisor_messages, last_supervisor_phase, sent_supervisor_event_count = (
                    extract_supervisor_stream_messages(
                        output,
                        last_phase=last_supervisor_phase,
                        sent_event_count=sent_supervisor_event_count,
                    )
                )
                for supervisor_message in supervisor_messages:
                    await websocket.send_json(supervisor_message)

                if HITLManager.is_paused(session_id) and event_type == "on_chain_start" and event_name == "agent_node":
                    await websocket.send_json({"type": "hitl_pause", "data": {"checkpoint_type": "checkpoint", "data": {}}})

                if event_type == "on_chat_model_start":
                    await websocket.send_json({"type": "thinking", "data": {"message": "Analyzing and reasoning..."}})
                elif event_type == "on_chat_model_end":
                    model_output = event.get("data", {}).get("output")
                    if model_output:
                        content = getattr(model_output, "content", "")
                        tool_calls = getattr(model_output, "tool_calls", [])
                        if tool_calls:
                            for tool_call in tool_calls:
                                await websocket.send_json(
                                    {
                                        "type": "tool_call",
                                        "data": {
                                            "tool": tool_call.get("name", ""),
                                            "input": tool_call.get("args", {}),
                                        },
                                    }
                                )
                        elif content:
                            output_scan = scan_llm_output(content)
                            await websocket.send_json(
                                {
                                    "type": "message",
                                    "data": {"content": output_scan.sanitized_content},
                                }
                            )
                elif event_type == "on_tool_end":
                    tool_output = event.get("data", {}).get("output", "")
                    content = str(tool_output.content) if hasattr(tool_output, "content") else str(tool_output)
                    await websocket.send_json(
                        {
                            "type": "tool_result",
                            "data": {"tool": event_name, "result": content[:2000]},
                        }
                    )
                    if event_name in ("write_todos", "get_todos"):
                        await websocket.send_json(
                            {
                                "type": "todo_update",
                                "data": {"todos": get_cached_session_todos(session_id)},
                            }
                        )
                    metrics = get_execution_metrics(session_id)
                    if metrics:
                        await websocket.send_json({"type": "metrics", "data": metrics})
                elif event_type == "on_chain_end" and event_name in {"Planning", "Replanning"}:
                    await websocket.send_json(
                        {
                            "type": "todo_update",
                            "data": {"todos": get_cached_session_todos(session_id)},
                        }
                    )
            except Exception as exc:
                logger.error("ws_event_error", error=str(exc), event_type=event.get("event", ""))

        if listener_task:
            listener_task.cancel()
        if heartbeat_task:
            heartbeat_task.cancel()

        metrics = get_execution_metrics(session_id)
        final_updates = {
            "status": "completed",
            "iterations_used": metrics.get("iterations_count", 0) if metrics else 0,
            "tokens_used": metrics.get("tokens_used", 0) if metrics else 0,
            "tool_calls_count": metrics.get("tool_calls_count", 0) if metrics else 0,
        }
        try:
            report = _workspace.read_file(session_id, "report.md")
            report, fabricated = verify_citations(report, tracked_urls)
            if fabricated:
                _workspace.write_file(session_id, "report.md", report)
            final_updates["final_report"] = report
        except FileNotFoundError:
            final_updates["final_report"] = ""

        session = await session_store.update(session_id, final_updates) or session
        await _update_session_in_db(session_id, str(user.id), final_updates)
        await audit_logger.log("session_complete", user_id=str(user.id), details={"session_id": session_id})
        await websocket.send_json({"type": "complete", "data": {"status": "completed", "session": session}})
    except WebSocketDisconnect:
        if user:
            await session_store.update(session_id, {"status": "cancelled"})
            await _update_session_in_db(session_id, str(user.id), {"status": "cancelled"})
        logger.info("ws_disconnected", session_id=session_id)
    except HTTPException as exc:
        if websocket.client_state.name != "CONNECTED":
            await websocket.close(code=4401, reason=exc.detail)
        else:
            await websocket.send_json({"type": "error", "data": {"message": exc.detail}})
            await websocket.close(code=4401)
    except Exception as exc:
        logger.error("ws_error", session_id=session_id, error=str(exc))
        if user:
            await session_store.update(session_id, {"status": "failed"})
            await _update_session_in_db(session_id, str(user.id), {"status": "failed"})
        try:
            await websocket.send_json({"type": "error", "data": {"message": str(exc)}})
        except Exception:
            pass
    finally:
        if listener_task:
            listener_task.cancel()
        if heartbeat_task:
            heartbeat_task.cancel()
        if tenant_token is not None:
            reset_tenant_context(tenant_token)
        cleanup_session(session_id)
