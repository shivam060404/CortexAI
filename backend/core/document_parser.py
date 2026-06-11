"""
Document Parser — unified file-to-text conversion layer.

Supports:
  - PDF   (via PyPDF2 / pypdf)
  - DOCX  (via python-docx)
  - Markdown (.md)
  - Plain text (.txt)
  - CSV
  - Images (base64 → vision LLM description via Groq)

Returns structured `ParsedDocument` with text, metadata, and optional
per-page/per-section content.
"""

from __future__ import annotations

import base64
import csv
import io
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from backend.core.logger import get_logger

logger = get_logger(__name__)

# Maximum upload size: 25 MB
MAX_UPLOAD_SIZE = 25 * 1024 * 1024

# Supported MIME types → extensions
SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
    ".csv": "csv",
    ".tsv": "csv",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".gif": "image",
    ".webp": "image",
    ".bmp": "image",
}

SUPPORTED_MIME_PREFIXES = [
    "application/pdf",
    "application/vnd.openxmlformats",
    "text/plain",
    "text/csv",
    "text/markdown",
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/webp",
    "image/bmp",
]


@dataclass
class ParsedDocument:
    """Result of parsing a single document."""
    filename: str
    file_type: str  # pdf, docx, markdown, text, csv, image
    text: str  # Full extracted text
    pages: list[str] = field(default_factory=list)  # Per-page/section content
    metadata: dict[str, Any] = field(default_factory=dict)
    word_count: int = 0
    parse_error: str = ""
    base64_data: str = ""  # For images: raw base64 for vision model

    @property
    def is_valid(self) -> bool:
        return bool(self.text.strip()) or bool(self.base64_data)

    def to_summary(self) -> str:
        """Short summary for logging / UI display."""
        status = "✓" if self.is_valid else "✗"
        return f"{status} {self.filename} ({self.file_type}, {self.word_count} words)"


class DocumentParser:
    """
    Unified document parser that converts uploaded files to text.

    Usage:
        parser = DocumentParser()
        doc = await parser.parse(filename="report.pdf", content=raw_bytes)
        print(doc.text)
    """

    def detect_type(self, filename: str, content_type: str = "") -> str:
        """Detect document type from filename extension or MIME type."""
        ext = Path(filename).suffix.lower()
        if ext in SUPPORTED_EXTENSIONS:
            return SUPPORTED_EXTENSIONS[ext]

        # Fallback: check MIME type
        for mime_prefix in SUPPORTED_MIME_PREFIXES:
            if content_type.startswith(mime_prefix):
                if "pdf" in mime_prefix:
                    return "pdf"
                elif "openxmlformats" in mime_prefix:
                    return "docx"
                elif "csv" in mime_prefix:
                    return "csv"
                elif "markdown" in mime_prefix:
                    return "markdown"
                elif "text" in mime_prefix:
                    return "text"
                elif "image" in mime_prefix:
                    return "image"

        return "unknown"

    async def parse(
        self,
        filename: str,
        content: bytes,
        content_type: str = "",
    ) -> ParsedDocument:
        """Parse a file and extract text content.

        Args:
            filename: Original filename (used for extension detection).
            content: Raw file bytes.
            content_type: Optional MIME type hint.

        Returns:
            ParsedDocument with extracted text and metadata.
        """
        if len(content) > MAX_UPLOAD_SIZE:
            return ParsedDocument(
                filename=filename,
                file_type="unknown",
                text="",
                parse_error=f"File too large ({len(content)} bytes, max {MAX_UPLOAD_SIZE})",
            )

        file_type = self.detect_type(filename, content_type)
        logger.info("parse_document_start", filename=filename, file_type=file_type, size=len(content))

        try:
            if file_type == "pdf":
                return self._parse_pdf(filename, content)
            elif file_type == "docx":
                return self._parse_docx(filename, content)
            elif file_type == "markdown":
                return self._parse_markdown(filename, content)
            elif file_type == "text":
                return self._parse_text(filename, content)
            elif file_type == "csv":
                return self._parse_csv(filename, content)
            elif file_type == "image":
                return self._parse_image(filename, content)
            else:
                return ParsedDocument(
                    filename=filename,
                    file_type="unknown",
                    text="",
                    parse_error=f"Unsupported file type: {Path(filename).suffix}",
                )
        except Exception as e:
            logger.error("parse_document_error", filename=filename, error=str(e))
            return ParsedDocument(
                filename=filename,
                file_type=file_type,
                text="",
                parse_error=str(e),
            )

    # ------------------------------------------------------------------
    # PDF Parser
    # ------------------------------------------------------------------
    def _parse_pdf(self, filename: str, content: bytes) -> ParsedDocument:
        """Extract text from PDF using PyPDF2."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            try:
                from pypdf import PdfReader
            except ImportError:
                return ParsedDocument(
                    filename=filename, file_type="pdf", text="",
                    parse_error="PyPDF2 not installed. Run: pip install PyPDF2",
                )

        reader = PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)

        full_text = "\n\n".join(pages)
        word_count = len(full_text.split())

        meta = {}
        if reader.metadata:
            meta = {
                "title": str(reader.metadata.get("/Title", "") or ""),
                "author": str(reader.metadata.get("/Author", "") or ""),
                "subject": str(reader.metadata.get("/Subject", "") or ""),
                "page_count": len(reader.pages),
            }

        return ParsedDocument(
            filename=filename,
            file_type="pdf",
            text=full_text.strip(),
            pages=pages,
            metadata={**meta, "page_count": len(reader.pages)},
            word_count=word_count,
        )

    # ------------------------------------------------------------------
    # DOCX Parser
    # ------------------------------------------------------------------
    def _parse_docx(self, filename: str, content: bytes) -> ParsedDocument:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            return ParsedDocument(
                filename=filename, file_type="docx", text="",
                parse_error="python-docx not installed. Run: pip install python-docx",
            )

        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name and "Heading" in para.style.name:
                    level = para.style.name.replace("Heading", "").strip()
                    try:
                        lvl = int(level)
                    except ValueError:
                        lvl = 2
                    text = f"{'#' * lvl} {text}"
                paragraphs.append(text)

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs)
        word_count = len(full_text.split())

        meta = {}
        if doc.core_properties:
            meta = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
            }

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            text=full_text.strip(),
            pages=paragraphs,
            metadata=meta,
            word_count=word_count,
        )

    # ------------------------------------------------------------------
    # Markdown Parser
    # ------------------------------------------------------------------
    def _parse_markdown(self, filename: str, content: bytes) -> ParsedDocument:
        """Parse markdown file, preserving structure as text."""
        text = content.decode("utf-8", errors="replace")
        word_count = len(text.split())

        # Split into sections by headings
        sections = re.split(r"^(#{1,3}\s+.+)$", text, flags=re.MULTILINE)

        return ParsedDocument(
            filename=filename,
            file_type="markdown",
            text=text.strip(),
            pages=[s.strip() for s in sections if s.strip()],
            metadata={"sections": len(sections)},
            word_count=word_count,
        )

    # ------------------------------------------------------------------
    # Plain Text Parser
    # ------------------------------------------------------------------
    def _parse_text(self, filename: str, content: bytes) -> ParsedDocument:
        """Parse plain text file."""
        # Try UTF-8, then latin-1
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1", errors="replace")

        word_count = len(text.split())
        # Split by double newlines into paragraphs
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        return ParsedDocument(
            filename=filename,
            file_type="text",
            text=text.strip(),
            pages=paragraphs,
            word_count=word_count,
        )

    # ------------------------------------------------------------------
    # CSV Parser
    # ------------------------------------------------------------------
    def _parse_csv(self, filename: str, content: bytes) -> ParsedDocument:
        """Parse CSV/TSV into structured text."""
        try:
            text = content.decode("utf-8", errors="replace")
        except UnicodeDecodeError:
            text = content.decode("latin-1", errors="replace")

        # Detect delimiter
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(text[:2048])
            delimiter = dialect.delimiter
        except csv.Error:
            delimiter = "\t" if filename.endswith(".tsv") else ","

        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = list(reader)

        if not rows:
            return ParsedDocument(
                filename=filename, file_type="csv", text="",
                word_count=0, parse_error="Empty CSV file",
            )

        # Format as readable text
        headers = rows[0] if rows else []
        lines = []

        if headers:
            lines.append("Columns: " + ", ".join(headers))
            lines.append("")

        for i, row in enumerate(rows[1:], 1):
            parts = []
            for h, v in zip(headers, row):
                if v.strip():
                    parts.append(f"{h}: {v.strip()}")
            if parts:
                lines.append(f"Row {i}: " + " | ".join(parts))

        full_text = "\n".join(lines)
        word_count = len(full_text.split())

        return ParsedDocument(
            filename=filename,
            file_type="csv",
            text=full_text.strip(),
            pages=[full_text],
            metadata={"rows": len(rows) - 1, "columns": headers},
            word_count=word_count,
        )

    # ------------------------------------------------------------------
    # Image Parser (base64 extraction for vision model)
    # ------------------------------------------------------------------
    def _parse_image(self, filename: str, content: bytes) -> ParsedDocument:
        """Encode image as base64 for downstream vision model analysis.

        The actual visual analysis happens in the vision_tools.analyze_image
        tool when the research agent invokes it. Here we just prepare the data.
        """
        ext = Path(filename).suffix.lower().lstrip(".")
        mime_map = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png",
                     "gif": "gif", "webp": "webp", "bmp": "bmp"}
        mime = mime_map.get(ext, "jpeg")

        b64 = base64.b64encode(content).decode("utf-8")
        data_uri = f"data:image/{mime};base64,{b64}"

        return ParsedDocument(
            filename=filename,
            file_type="image",
            text=f"[Image: {filename}] — attached for visual analysis",
            base64_data=data_uri,
            metadata={"mime_type": f"image/{mime}", "size_bytes": len(content)},
            word_count=0,
        )


# ---------------------------------------------------------------------------
# Batch parsing for multiple files
# ---------------------------------------------------------------------------

async def parse_documents(
    files: list[tuple[str, bytes, str]],
) -> list[ParsedDocument]:
    """Parse multiple files in batch.

    Args:
        files: List of (filename, content_bytes, content_type) tuples.

    Returns:
        List of ParsedDocument results.
    """
    parser = DocumentParser()
    results = []
    for filename, content, content_type in files:
        doc = await parser.parse(filename, content, content_type)
        results.append(doc)
        logger.info("batch_parse_result", summary=doc.to_summary())
    return results


# Module-level singleton
document_parser = DocumentParser()
